#!/usr/bin/env python3
"""
Enhanced Robust AI Feedback Generator
Handles complex documents with text, code, figures, tables, and any structure.

Features:
- Multi-format document processing (PDF, DOCX, ZIP, TXT, code files)
- Advanced text extraction with OCR fallback
- Table and figure detection
- Code analysis capabilities
- Structured document analysis
- Robust error handling and fallbacks
"""

import os
import logging
from typing import Optional, List, Dict, Any, Union
from pathlib import Path
import tempfile
import zipfile
import re
import io
from dotenv import load_dotenv
import base64

# Load environment variables
load_dotenv()

# Enhanced document processing imports
import PyPDF2
from docx import Document
from docx.document import Document as DocumentType
import pandas as pd

# Advanced PDF processing with multiple engines
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Claude API (always available if anthropic package is installed)
try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# RAG system imports (optional — requires sentence-transformers + faiss which pull torch)
try:
    from sentence_transformers import SentenceTransformer
    import faiss
    import numpy as np
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    """Enhanced document processor with multiple extraction engines."""

    @staticmethod
    def extract_text_from_pdf_advanced(file_path: str) -> Dict[str, Any]:
        """Advanced PDF extraction with multiple methods and structure analysis."""
        result = {
            'text': '',
            'tables': [],
            'images_found': 0,
            'code_blocks': [],
            'structure': {},
            'extraction_method': '',
            'success': False
        }

        # Method 1: Try pdfplumber (best for tables and structure)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    tables = []

                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"=== PAGE {page_num + 1} ===\n{page_text}\n")

                        # Extract tables
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table_idx, table in enumerate(page_tables):
                                table_text = f"\n=== TABLE {table_idx + 1} ON PAGE {page_num + 1} ===\n"
                                for row in table:
                                    if row:
                                        table_text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                                tables.append(table_text)
                                text_parts.append(table_text)

                    result['text'] = "\n".join(text_parts)
                    result['tables'] = tables
                    result['extraction_method'] = 'pdfplumber'
                    result['success'] = True
                    logger.info(f"PDF extracted with pdfplumber: {len(result['text'])} chars, {len(tables)} tables")
                    return result

            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")

        # Method 2: Try PyMuPDF (good for images and complex layouts)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text_parts = []
                images_found = 0

                for page_num in range(doc.page_count):
                    page = doc[page_num]

                    # Extract text
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(f"=== PAGE {page_num + 1} ===\n{page_text}\n")

                    # Check for images
                    image_list = page.get_images()
                    images_found += len(image_list)

                    if image_list:
                        text_parts.append(f"[PAGE {page_num + 1} CONTAINS {len(image_list)} IMAGES/FIGURES]\n")

                doc.close()
                result['text'] = "\n".join(text_parts)
                result['images_found'] = images_found
                result['extraction_method'] = 'PyMuPDF'
                result['success'] = True
                logger.info(f"PDF extracted with PyMuPDF: {len(result['text'])} chars, {images_found} images")
                return result

            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")

        # Method 3: Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"=== PAGE {page_num + 1} ===\n{page_text}\n")

                result['text'] = "\n".join(text_parts)
                result['extraction_method'] = 'PyPDF2'
                result['success'] = True
                logger.info(f"PDF extracted with PyPDF2: {len(result['text'])} chars")
                return result

        except Exception as e:
            logger.error(f"All PDF extraction methods failed: {e}")

        return result

    @staticmethod
    def extract_text_from_docx_advanced(file_path: str) -> Dict[str, Any]:
        """Advanced DOCX extraction with structure analysis."""
        result = {
            'text': '',
            'tables': [],
            'images_found': 0,
            'code_blocks': [],
            'structure': {},
            'success': False
        }

        try:
            doc = Document(file_path)
            text_parts = []
            tables = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = f"\n=== TABLE {table_idx + 1} ===\n"
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_text += " | ".join(row_data) + "\n"
                tables.append(table_text)
                text_parts.append(table_text)

            # Count images (rough estimate)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    result['images_found'] += 1

            result['text'] = "\n".join(text_parts)
            result['tables'] = tables
            result['success'] = True
            logger.info(f"DOCX extracted: {len(result['text'])} chars, {len(tables)} tables, ~{result['images_found']} images")

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")

        return result

    @staticmethod
    def extract_text_from_code_file(file_path: str, filename: str) -> Dict[str, Any]:
        """Extract and analyze code files with structure detection."""
        result = {
            'text': '',
            'code_blocks': [],
            'structure': {},
            'language': '',
            'success': False
        }

        try:
            # Detect language from extension
            ext = Path(filename).suffix.lower()
            language_map = {
                '.py': 'Python', '.java': 'Java', '.cpp': 'C++', '.c': 'C',
                '.js': 'JavaScript', '.ts': 'TypeScript', '.html': 'HTML',
                '.css': 'CSS', '.sql': 'SQL', '.r': 'R', '.m': 'MATLAB',
                '.php': 'PHP', '.rb': 'Ruby', '.go': 'Go', '.rs': 'Rust'
            }

            result['language'] = language_map.get(ext, 'Text')

            # Read file with multiple encoding attempts
            encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise ValueError("Could not decode file with any encoding")

            # Analyze code structure
            lines = content.split('\n')
            structure = {
                'total_lines': len(lines),
                'functions': [],
                'classes': [],
                'imports': [],
                'comments': 0
            }

            # Basic code analysis
            for i, line in enumerate(lines, 1):
                line_stripped = line.strip()

                # Count comments
                if line_stripped.startswith('#') or line_stripped.startswith('//') or line_stripped.startswith('/*'):
                    structure['comments'] += 1

                # Detect functions (basic patterns)
                if any(pattern in line_stripped for pattern in ['def ', 'function ', 'public ', 'private ', 'protected ']):
                    structure['functions'].append(f"Line {i}: {line_stripped[:50]}...")

                # Detect classes
                if any(pattern in line_stripped for pattern in ['class ', 'interface ', 'struct ']):
                    structure['classes'].append(f"Line {i}: {line_stripped[:50]}...")

                # Detect imports
                if any(pattern in line_stripped for pattern in ['import ', 'include ', 'using ', 'require']):
                    structure['imports'].append(line_stripped)

            # Format the text with structure information
            text_parts = [
                f"=== {result['language']} CODE FILE: {filename} ===",
                f"Total Lines: {structure['total_lines']}",
                f"Functions Found: {len(structure['functions'])}",
                f"Classes Found: {len(structure['classes'])}",
                f"Import Statements: {len(structure['imports'])}",
                f"Comment Lines: {structure['comments']}",
                "",
                "=== CODE CONTENT ===",
                content
            ]

            if structure['functions']:
                text_parts.extend(["", "=== FUNCTIONS DETECTED ==="] + structure['functions'][:10])

            if structure['classes']:
                text_parts.extend(["", "=== CLASSES DETECTED ==="] + structure['classes'][:10])

            result['text'] = "\n".join(text_parts)
            result['code_blocks'] = [content]
            result['structure'] = structure
            result['success'] = True

            logger.info(f"Code file analyzed: {result['language']}, {structure['total_lines']} lines")

        except Exception as e:
            logger.error(f"Code file extraction failed: {e}")

        return result

    @staticmethod
    def extract_text_from_zip_advanced(file_path: str) -> Dict[str, Any]:
        """Advanced ZIP extraction with file type analysis."""
        result = {
            'text': '',
            'files_processed': 0,
            'file_types': {},
            'structure': {},
            'success': False
        }

        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                text_parts = [f"=== ZIP ARCHIVE ANALYSIS ==="]
                text_parts.append(f"Total Files: {len(file_list)}")

                # Analyze file types
                file_types = {}
                for file_name in file_list:
                    if not file_name.endswith('/'):  # Skip directories
                        ext = Path(file_name).suffix.lower()
                        file_types[ext] = file_types.get(ext, 0) + 1

                result['file_types'] = file_types
                text_parts.append(f"File Types: {dict(file_types)}")
                text_parts.append("")

                # Process each file
                for file_name in file_list[:20]:  # Limit to first 20 files
                    if file_name.endswith('/'):
                        continue

                    try:
                        with zip_file.open(file_name) as file_content:
                            ext = Path(file_name).suffix.lower()

                            text_parts.append(f"=== FILE: {file_name} ===")

                            if ext in ['.txt', '.md', '.py', '.java', '.cpp', '.c', '.js', '.html', '.css']:
                                # Text-based files
                                try:
                                    content = file_content.read().decode('utf-8')
                                    text_parts.append(content[:2000])  # Limit content
                                    if len(content) > 2000:
                                        text_parts.append("... [CONTENT TRUNCATED] ...")
                                except UnicodeDecodeError:
                                    text_parts.append("[BINARY OR ENCODED CONTENT - CANNOT DISPLAY]")

                            elif ext in ['.pdf', '.docx']:
                                text_parts.append(f"[{ext.upper()} FILE DETECTED - REQUIRES SEPARATE PROCESSING]")

                            else:
                                file_size = len(file_content.read())
                                text_parts.append(f"[BINARY FILE - SIZE: {file_size} bytes]")

                            result['files_processed'] += 1

                    except Exception as e:
                        text_parts.append(f"[ERROR PROCESSING FILE: {e}]")

                result['text'] = "\n".join(text_parts)
                result['success'] = True
                logger.info(f"ZIP processed: {result['files_processed']} files, types: {file_types}")

        except Exception as e:
            logger.error(f"ZIP extraction failed: {e}")

        return result

class EnhancedFeedbackGenerator:
    """Enhanced feedback generator with robust document processing."""

    def __init__(self):
        """Initialize the enhanced feedback generator."""
        self.doc_processor = EnhancedDocumentProcessor()
        self.rag_system = None
        self.claude_client = None
        self.embedding_model = None
        self.index = None
        self.chunks = []

        # Initialize Claude client independently of RAG dependencies
        if ANTHROPIC_AVAILABLE:
            api_key = os.getenv('ANTHROPIC_API_KEY')
            if api_key:
                try:
                    self.claude_client = anthropic.Anthropic(api_key=api_key)
                    logger.info("Claude client initialized successfully")
                except Exception as e:
                    logger.error(f"Failed to initialize Claude client: {e}")
            else:
                logger.warning("No ANTHROPIC_API_KEY found — AI feedback will be unavailable")

        # Optional: RAG with embeddings (requires sentence-transformers + faiss)
        if RAG_AVAILABLE:
            try:
                self.embedding_model = SentenceTransformer('BAAI/bge-small-en-v1.5')
                self.index = faiss.IndexFlatIP(384)
                logger.info("RAG embedding model loaded")
            except Exception as e:
                logger.warning(f"RAG model load failed (non-fatal): {e}")

    def generate_enhanced_feedback(self, submission_file, rubric_file=None) -> str:
        """Generate enhanced feedback with robust document processing."""
        try:
            # Extract text with advanced processing
            extraction_result = self._extract_text_advanced(submission_file)

            if not extraction_result['success'] or not extraction_result['text'].strip():
                return "Error: Could not extract any readable content from the submission file. The file may be corrupted, encrypted, or in an unsupported format."

            # Build comprehensive document analysis
            document_analysis = self._build_document_analysis(extraction_result)

            # Extract rubric if provided
            rubric_text = ""
            if rubric_file:
                rubric_result = self._extract_text_advanced(rubric_file)
                if rubric_result['success']:
                    rubric_text = rubric_result['text']

            # Generate feedback using Claude
            feedback = self._generate_claude_feedback(document_analysis, rubric_text)

            return feedback

        except Exception as e:
            logger.error(f"Enhanced feedback generation failed: {e}")
            return f"Error generating feedback: {str(e)}. Please ensure the file is not corrupted and try again."

    def _extract_text_advanced(self, file_obj) -> Dict[str, Any]:
        """Advanced text extraction with multiple fallback methods."""
        temp_file_path = None
        try:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file_obj.name).suffix) as temp_file:
                temp_file_path = temp_file.name
                # Handle both Streamlit UploadedFile and regular file objects
                if hasattr(file_obj, 'getvalue'):
                    temp_file.write(file_obj.getvalue())
                elif hasattr(file_obj, 'read'):
                    temp_file.write(file_obj.read())
                else:
                    raise ValueError(f"Unsupported file object type: {type(file_obj)}")
                temp_file.flush()

            # Process the file after closing the handle
            file_extension = Path(file_obj.name).suffix.lower()
            logger.info(f"Processing file: {file_obj.name} ({file_extension})")

            # Route to appropriate extraction method
            if file_extension == '.pdf':
                result = self.doc_processor.extract_text_from_pdf_advanced(temp_file_path)
            elif file_extension == '.docx':
                result = self.doc_processor.extract_text_from_docx_advanced(temp_file_path)
            elif file_extension == '.zip':
                result = self.doc_processor.extract_text_from_zip_advanced(temp_file_path)
            elif file_extension in ['.py', '.java', '.cpp', '.c', '.js', '.ts', '.html', '.css', '.sql', '.r', '.m', '.php', '.rb', '.go', '.rs']:
                result = self.doc_processor.extract_text_from_code_file(temp_file_path, file_obj.name)
            else:
                # Text files and other formats
                result = self._extract_text_file_robust(temp_file_path, file_obj.name)

            return result

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return {'success': False, 'text': '', 'error': str(e)}
        finally:
            # Ensure cleanup happens
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temp file: {cleanup_error}")

    def _extract_text_file_robust(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Robust text file extraction with multiple encoding support."""
        result = {'text': '', 'success': False}

        encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252', 'ascii']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()

                result['text'] = f"=== TEXT FILE: {filename} ===\n{content}"
                result['success'] = True
                logger.info(f"Text file read with {encoding} encoding: {len(content)} chars")
                return result

            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"Failed to read with {encoding}: {e}")
                continue

        # If all encodings fail, try binary mode and decode what we can
        try:
            with open(file_path, 'rb') as f:
                binary_content = f.read()

            # Try to decode as much as possible
            decoded_content = binary_content.decode('utf-8', errors='ignore')
            result['text'] = f"=== BINARY FILE (PARTIAL DECODE): {filename} ===\n{decoded_content}"
            result['success'] = True
            logger.info(f"Binary file partially decoded: {len(decoded_content)} chars")

        except Exception as e:
            logger.error(f"All text extraction methods failed: {e}")

        return result

    def _build_document_analysis(self, extraction_result: Dict[str, Any]) -> str:
        """Build comprehensive document analysis."""
        analysis_parts = ["=== COMPREHENSIVE DOCUMENT ANALYSIS ===\n"]

        # Basic document info
        analysis_parts.append(f"Document Processing: {extraction_result.get('extraction_method', 'Unknown')}")
        analysis_parts.append(f"Content Length: {len(extraction_result.get('text', ''))} characters")

        # Structure analysis
        if extraction_result.get('tables'):
            analysis_parts.append(f"Tables Found: {len(extraction_result['tables'])}")

        if extraction_result.get('images_found', 0) > 0:
            analysis_parts.append(f"Images/Figures Found: {extraction_result['images_found']}")

        if extraction_result.get('code_blocks'):
            analysis_parts.append(f"Code Blocks: {len(extraction_result['code_blocks'])}")

        if extraction_result.get('language'):
            analysis_parts.append(f"Programming Language: {extraction_result['language']}")

        if extraction_result.get('structure'):
            struct = extraction_result['structure']
            if isinstance(struct, dict):
                for key, value in struct.items():
                    if isinstance(value, (int, str)):
                        analysis_parts.append(f"{key.replace('_', ' ').title()}: {value}")

        analysis_parts.append("\n=== DOCUMENT CONTENT ===")
        analysis_parts.append(extraction_result.get('text', ''))

        return "\n".join(analysis_parts)

    def _generate_claude_feedback(self, document_analysis: str, rubric_text: str = "") -> str:
        """Generate feedback using Claude API with enhanced prompting."""
        try:
            if not self.claude_client:
                return "Error: Claude API not configured. Please set ANTHROPIC_API_KEY in environment variables."

            # Build comprehensive prompt
            prompt_parts = [
                "You are an expert academic assessor providing detailed feedback on student submissions.",
                "Please provide comprehensive, constructive feedback that covers:",
                "1. Content quality and accuracy",
                "2. Structure and organization",
                "3. Technical implementation (if applicable)",
                "4. Use of figures, tables, and visual elements",
                "5. Code quality and logic (if applicable)",
                "6. Areas of strength",
                "7. Areas for improvement",
                "8. Specific recommendations",
                "",
                "Document Analysis:"
            ]

            if rubric_text.strip():
                prompt_parts.extend([
                    "",
                    "MARKING RUBRIC:",
                    rubric_text,
                    "",
                    "Please assess the submission against this rubric."
                ])

            prompt = "\n".join(prompt_parts)

            # Generate feedback
            message = self.claude_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=2000,
                messages=[
                    {
                        "role": "user",
                        "content": f"{prompt}\n\n{document_analysis}"
                    }
                ]
            )

            feedback = message.content[0].text

            # Add system footer
            footer = "\n\n--- AI-Generated Feedback (Claude API - For Reference Only) ---\nThis feedback was generated by Claude AI to support the marking process. Human markers should use their professional judgment for final assessment decisions."

            return feedback + footer

        except Exception as e:
            logger.error(f"Claude feedback generation failed: {e}")
            return f"Error generating AI feedback: {str(e)}. Please check your API configuration and try again."

# Global instance
enhanced_feedback_generator = EnhancedFeedbackGenerator()