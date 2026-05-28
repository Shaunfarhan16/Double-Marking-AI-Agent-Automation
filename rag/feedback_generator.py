import os
import logging
from typing import Optional, List, Dict, Any
from pathlib import Path
import tempfile
import zipfile
import re
import io
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Document processing imports with multiple fallbacks
import PyPDF2
from docx import Document
import pandas as pd

# Advanced PDF processing imports
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# RAG system imports
try:
    import anthropic
    from sentence_transformers import SentenceTransformer
    import faiss
    import numpy as np
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False
    logging.warning("RAG dependencies not available. Install anthropic, sentence-transformers, faiss-cpu")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processor with multiple PDF engines and comprehensive content extraction."""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Multi-engine PDF text extraction with figure, table, and code detection."""

        # Verify file accessibility
        if not DocumentProcessor._verify_file(file_path):
            return ""

        # Try multiple extraction methods in order of sophistication
        extraction_methods = [
            DocumentProcessor._extract_with_pdfplumber,
            DocumentProcessor._extract_with_pymupdf,
            DocumentProcessor._extract_with_pypdf2,
        ]

        text = ""
        method_used = "none"

        for method in extraction_methods:
            try:
                logger.info(f"Trying extraction method: {method.__name__}")
                text = method(file_path)
                method_used = method.__name__
                logger.info(f"Method {method.__name__} returned {len(text) if text else 0} characters")
                if text and text.strip():
                    logger.info(f"Successfully extracted text using {method_used}")
                    break
                else:
                    logger.warning(f"Method {method.__name__} returned empty text")
            except Exception as e:
                logger.error(f"Method {method.__name__} failed with error: {e}")
                import traceback
                logger.error(f"Full traceback: {traceback.format_exc()}")
                continue

        # If no text extracted, try OCR as last resort
        if not text.strip() and OCR_AVAILABLE:
            logger.info("Attempting OCR extraction for scanned PDF")
            text = DocumentProcessor._extract_with_ocr(file_path)
            method_used = "OCR"

        if not text.strip():
            logger.error("All PDF extraction methods failed")
            return ""

        # Enhanced content analysis and structuring
        structured_text = DocumentProcessor._analyze_and_structure_content(text, file_path, method_used)
        logger.info(f"Final extraction: {len(structured_text)} characters using {method_used}")

        return structured_text

    @staticmethod
    def _verify_file(file_path: str) -> bool:
        """Verify file exists and is accessible."""
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return False

            file_size = os.path.getsize(file_path) / (1024 * 1024)  # MB
            if file_size > 100:  # Increased limit to 100MB
                logger.error(f"File too large: {file_size:.2f}MB")
                return False

            # Test if file can be opened
            with open(file_path, 'rb') as f:
                f.read(10)  # Read first 10 bytes

            return True
        except Exception as e:
            logger.error(f"File verification failed: {e}")
            return False

    @staticmethod
    def _extract_with_pdfplumber(file_path: str) -> str:
        """Extract text using pdfplumber (best for tables and structured content)."""
        if not PDFPLUMBER_AVAILABLE:
            logger.error("pdfplumber is not available")
            raise ImportError("pdfplumber not available")

        logger.info("Starting pdfplumber extraction")
        text = ""

        try:
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"pdfplumber opened PDF with {len(pdf.pages)} pages")

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        # Extract text
                        page_text = page.extract_text() or ""
                        logger.info(f"pdfplumber page {page_num}: {len(page_text)} characters of text")

                        # Extract tables
                        tables = page.extract_tables()
                        table_text = ""
                        if tables:
                            logger.info(f"pdfplumber page {page_num}: found {len(tables)} tables")
                            for i, table in enumerate(tables):
                                table_text += f"\n=== TABLE {i+1} ON PAGE {page_num} ===\n"
                                for row in table:
                                    if row:
                                        table_text += " | ".join([str(cell or "") for cell in row]) + "\n"

                        # Combine text and tables
                        page_content = f"\n=== PAGE {page_num} ===\n{page_text}\n{table_text}"
                        text += page_content

                    except Exception as e:
                        logger.error(f"pdfplumber failed on page {page_num}: {e}")
                        continue

                logger.info(f"pdfplumber total extraction: {len(text)} characters")

        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")
            raise

        return text

    @staticmethod
    def _extract_with_pymupdf(file_path: str) -> str:
        """Extract text using PyMuPDF (good for complex layouts)."""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available")

        text = ""
        pdf_document = fitz.open(file_path)

        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]

            # Extract text
            page_text = page.get_text()

            # Extract images and figures info
            image_list = page.get_images()
            image_info = ""
            if image_list:
                image_info = f"\n=== FIGURES/IMAGES ON PAGE {page_num + 1} ===\n"
                for img_index, img in enumerate(image_list):
                    image_info += f"Figure {img_index + 1}: Image detected (xref: {img[0]})\n"

            # Extract tables (basic detection)
            tables = page.find_tables()
            table_info = ""
            if tables:
                table_info = f"\n=== TABLES ON PAGE {page_num + 1} ===\n"
                for table_index, table in enumerate(tables):
                    table_info += f"Table {table_index + 1}: {table.bbox}\n"
                    try:
                        table_data = table.extract()
                        for row in table_data:
                            table_info += " | ".join([str(cell or "") for cell in row]) + "\n"
                    except:
                        table_info += "Table data extraction failed\n"

            page_content = f"\n=== PAGE {page_num + 1} ===\n{page_text}\n{image_info}{table_info}"
            text += page_content

        pdf_document.close()
        return text

    @staticmethod
    def _extract_with_pypdf2(file_path: str) -> str:
        """Extract text using PyPDF2 (fallback method)."""
        text = ""

        try:
            with open(file_path, 'rb') as file:
                logger.info("Opening PDF with PyPDF2")
                pdf_reader = PyPDF2.PdfReader(file)

                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.error("PDF is encrypted - cannot process")
                    return ""

                num_pages = len(pdf_reader.pages)
                logger.info(f"PDF has {num_pages} pages")

                if num_pages == 0:
                    logger.error("PDF has no pages")
                    return ""

                # Extract text from each page
                pages_with_text = 0
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        logger.info(f"Page {page_num}: extracted {len(page_text) if page_text else 0} characters")

                        if page_text and page_text.strip():
                            text += f"\n=== PAGE {page_num} ===\n{page_text}\n"
                            pages_with_text += 1
                        else:
                            logger.warning(f"Page {page_num} has no extractable text")

                    except Exception as e:
                        logger.error(f"Page {page_num} extraction failed: {e}")
                        continue

                logger.info(f"PyPDF2: {pages_with_text}/{num_pages} pages had extractable text")

        except Exception as e:
            logger.error(f"PyPDF2 extraction completely failed: {e}")
            return ""

        return text

    @staticmethod
    def _extract_with_ocr(file_path: str) -> str:
        """Extract text using OCR for scanned PDFs."""
        if not OCR_AVAILABLE:
            raise ImportError("OCR dependencies not available")

        text = ""
        try:
            # Convert PDF to images and apply OCR
            pdf_document = fitz.open(file_path)

            for page_num in range(min(pdf_document.page_count, 10)):  # Limit to 10 pages for OCR
                page = pdf_document[page_num]

                # Convert page to image
                mat = fitz.Matrix(2, 2)  # Increase resolution
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")

                # Apply OCR
                image = Image.open(io.BytesIO(img_data))
                page_text = pytesseract.image_to_string(image)

                if page_text.strip():
                    text += f"\n=== PAGE {page_num + 1} (OCR) ===\n{page_text}\n"

            pdf_document.close()

        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""

        return text

    @staticmethod
    def _analyze_and_structure_content(text: str, file_path: str, method_used: str) -> str:
        """Analyze and structure extracted content for better RAG processing."""

        structured_content = f"=== DOCUMENT ANALYSIS ===\n"
        structured_content += f"File: {Path(file_path).name}\n"
        structured_content += f"Extraction Method: {method_used}\n"
        structured_content += f"Content Length: {len(text)} characters\n\n"

        # Analyze content types
        content_analysis = DocumentProcessor._analyze_content_types(text)

        structured_content += "=== CONTENT ANALYSIS ===\n"
        for content_type, details in content_analysis.items():
            if details['count'] > 0:
                structured_content += f"{content_type}: {details['count']} found\n"
                if details['examples']:
                    structured_content += f"  Examples: {', '.join(details['examples'][:3])}\n"

        structured_content += "\n=== DOCUMENT SECTIONS ===\n"
        sections = DocumentProcessor._extract_sections(text)
        for section in sections:
            structured_content += f"- {section}\n"

        structured_content += "\n=== FULL CONTENT ===\n"
        structured_content += text

        return structured_content

    @staticmethod
    def _analyze_content_types(text: str) -> Dict:
        """Analyze different types of content in the document."""
        analysis = {
            'Academic Sections': {'count': 0, 'examples': []},
            'Figures': {'count': 0, 'examples': []},
            'Tables': {'count': 0, 'examples': []},
            'Code Blocks': {'count': 0, 'examples': []},
            'Mathematical Content': {'count': 0, 'examples': []},
            'References': {'count': 0, 'examples': []}
        }

        text_lower = text.lower()

        # Academic sections
        section_patterns = [
            r'abstract', r'introduction', r'literature review', r'methodology',
            r'method', r'implementation', r'results', r'discussion',
            r'conclusion', r'references', r'appendix'
        ]
        for pattern in section_patterns:
            matches = re.findall(rf'\b{pattern}\b', text_lower)
            if matches:
                analysis['Academic Sections']['count'] += len(matches)
                analysis['Academic Sections']['examples'].extend(matches[:2])

        # Figures and images
        figure_patterns = [
            r'figure \d+', r'fig\. \d+', r'image \d+', r'chart \d+',
            r'graph \d+', r'diagram \d+', r'flowchart \d+'
        ]
        for pattern in figure_patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                analysis['Figures']['count'] += len(matches)
                analysis['Figures']['examples'].extend(matches[:3])

        # Tables
        table_patterns = [r'table \d+', r'tab\. \d+', r'=== table', r'\|.*\|']
        for pattern in table_patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                analysis['Tables']['count'] += len(matches)
                analysis['Tables']['examples'].extend(matches[:3])

        # Code blocks
        code_patterns = [
            r'def \w+', r'class \w+', r'function \w+', r'import \w+',
            r'#include', r'public class', r'console\.log', r'algorithm \d+'
        ]
        for pattern in code_patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                analysis['Code Blocks']['count'] += len(matches)
                analysis['Code Blocks']['examples'].extend(matches[:3])

        # Mathematical content
        math_patterns = [r'equation \d+', r'formula \d+', r'\$.*\$', r'\\begin\{']
        for pattern in math_patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                analysis['Mathematical Content']['count'] += len(matches)
                analysis['Mathematical Content']['examples'].extend(matches[:3])

        # References
        ref_patterns = [r'\[\d+\]', r'\(\d{4}\)', r'et al\.', r'doi:']
        for pattern in ref_patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                analysis['References']['count'] += len(matches)
                analysis['References']['examples'].extend(matches[:3])

        return analysis

    @staticmethod
    def _extract_sections(text: str) -> List[str]:
        """Extract document sections and headings."""
        sections = []

        # Look for numbered sections
        section_pattern = r'^(\d+\.?\s+[A-Z][A-Za-z\s]+)$'
        for line in text.split('\n'):
            line = line.strip()
            if re.match(section_pattern, line):
                sections.append(line)

        # Look for common academic sections
        academic_sections = [
            'abstract', 'introduction', 'literature review', 'methodology',
            'method', 'implementation', 'results', 'discussion', 'conclusion',
            'references', 'appendix'
        ]

        for line in text.split('\n'):
            line_lower = line.strip().lower()
            for section in academic_sections:
                if line_lower == section or line_lower.startswith(f"{section} "):
                    sections.append(line.strip())

        return list(set(sections))  # Remove duplicates

    @staticmethod
    def _analyze_page_content(page_text: str, page_num: int) -> Dict:
        """Analyze page content for figures, tables, sections, and code."""
        page_analysis = {
            'page_num': page_num,
            'content_type': 'text',
            'sections': [],
            'figures': [],
            'tables': [],
            'code_blocks': [],
            'has_appendix': False
        }

        text_lower = page_text.lower()

        # Section detection (enhanced)
        section_patterns = [
            r'^\s*(\d+\.?\s*)?abstract\s*$',
            r'^\s*(\d+\.?\s*)?introduction\s*$',
            r'^\s*(\d+\.?\s*)?literature review\s*$',
            r'^\s*(\d+\.?\s*)?methodology\s*$',
            r'^\s*(\d+\.?\s*)?method\s*$',
            r'^\s*(\d+\.?\s*)?implementation\s*$',
            r'^\s*(\d+\.?\s*)?results\s*$',
            r'^\s*(\d+\.?\s*)?discussion\s*$',
            r'^\s*(\d+\.?\s*)?conclusion\s*$',
            r'^\s*(\d+\.?\s*)?references\s*$',
            r'^\s*(\d+\.?\s*)?appendix\s*[a-z]?\s*$'
        ]

        import re
        for line in page_text.split('\n'):
            for pattern in section_patterns:
                if re.match(pattern, line.strip(), re.IGNORECASE):
                    page_analysis['sections'].append(line.strip())

        # Figure detection
        figure_indicators = [
            'figure', 'fig.', 'fig ', 'chart', 'graph', 'diagram',
            'flowchart', 'image', 'plot', 'visualization'
        ]

        for indicator in figure_indicators:
            if indicator in text_lower:
                # Find figure references
                figure_matches = re.findall(rf'{indicator}[\s]*[\d\w\.]*[\s]*[:;]?[\s]*([^\n]*)', text_lower)
                page_analysis['figures'].extend(figure_matches[:3])  # Limit to avoid noise

        # Table detection
        table_indicators = ['table', 'tab.', 'tab ']
        for indicator in table_indicators:
            if indicator in text_lower:
                table_matches = re.findall(rf'{indicator}[\s]*[\d\w\.]*[\s]*[:;]?[\s]*([^\n]*)', text_lower)
                page_analysis['tables'].extend(table_matches[:3])

        # Code block detection
        code_indicators = [
            'def ', 'class ', 'function', 'import ', 'from ', '#include',
            'public class', 'private ', 'void ', 'int main', 'console.log',
            '```', 'code:', 'listing', 'algorithm'
        ]

        for indicator in code_indicators:
            if indicator in text_lower:
                page_analysis['code_blocks'].append(f"Code detected: {indicator}")
                page_analysis['content_type'] = 'code'

        # Appendix detection
        if any(word in text_lower for word in ['appendix', 'annex', 'attachment']):
            page_analysis['has_appendix'] = True
            page_analysis['content_type'] = 'appendix'

        return page_analysis

    @staticmethod
    def _structure_academic_report(raw_text: str, structure: Dict) -> str:
        """Structure the academic report text for better RAG processing."""
        structured_content = "=== ACADEMIC REPORT ANALYSIS ===\n\n"

        # Document overview
        total_pages = len(structure['pages'])
        sections_found = []
        figures_count = 0
        tables_count = 0
        code_pages = 0
        appendix_pages = 0

        for page in structure['pages']:
            sections_found.extend(page['sections'])
            figures_count += len(page['figures'])
            tables_count += len(page['tables'])
            if page['content_type'] == 'code':
                code_pages += 1
            if page['has_appendix']:
                appendix_pages += 1

        # Add structure summary
        structured_content += f"DOCUMENT STRUCTURE:\n"
        structured_content += f"- Total Pages: {total_pages}\n"
        structured_content += f"- Sections: {len(set(sections_found))}\n"
        structured_content += f"- Figures/Charts: ~{figures_count}\n"
        structured_content += f"- Tables: ~{tables_count}\n"
        structured_content += f"- Code/Appendix Pages: {code_pages + appendix_pages}\n\n"

        if sections_found:
            structured_content += f"IDENTIFIED SECTIONS:\n"
            for section in set(sections_found):
                structured_content += f"- {section}\n"
            structured_content += "\n"

        # Add visual elements summary
        all_figures = []
        all_tables = []
        for page in structure['pages']:
            all_figures.extend(page['figures'])
            all_tables.extend(page['tables'])

        if all_figures:
            structured_content += "FIGURES & VISUAL ELEMENTS:\n"
            for i, fig in enumerate(set(all_figures)[:10], 1):
                if fig.strip():
                    structured_content += f"- Figure {i}: {fig[:100]}\n"
            structured_content += "\n"

        if all_tables:
            structured_content += "TABLES & DATA:\n"
            for i, table in enumerate(set(all_tables)[:5], 1):
                if table.strip():
                    structured_content += f"- Table {i}: {table[:100]}\n"
            structured_content += "\n"

        # Add original content
        structured_content += "=== FULL DOCUMENT CONTENT ===\n"
        structured_content += raw_text

        return structured_content
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Enhanced DOCX extraction with table and structure detection."""
        try:
            if not DocumentProcessor._verify_file(file_path):
                return ""

            doc = Document(file_path)
            text_content = ""

            # Extract main document text
            document_text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    document_text += paragraph.text + "\n"

            # Extract tables
            table_content = ""
            for table_num, table in enumerate(doc.tables, 1):
                table_content += f"\n=== TABLE {table_num} ===\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip() or "")
                    table_content += " | ".join(row_text) + "\n"

            # Extract headers and footers
            header_footer_content = ""
            try:
                for section in doc.sections:
                    if section.header:
                        header_text = ""
                        for paragraph in section.header.paragraphs:
                            header_text += paragraph.text + " "
                        if header_text.strip():
                            header_footer_content += f"HEADER: {header_text.strip()}\n"

                    if section.footer:
                        footer_text = ""
                        for paragraph in section.footer.paragraphs:
                            footer_text += paragraph.text + " "
                        if footer_text.strip():
                            header_footer_content += f"FOOTER: {footer_text.strip()}\n"
            except:
                pass  # Headers/footers are optional

            # Combine all content
            full_text = document_text + table_content + header_footer_content

            if not full_text.strip():
                logger.error("No text extracted from DOCX file")
                return ""

            # Analyze and structure content
            structured_text = DocumentProcessor._analyze_and_structure_content(
                full_text, file_path, "DOCX"
            )

            logger.info(f"Successfully extracted {len(structured_text)} characters from DOCX")
            return structured_text

        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_zip(file_path: str) -> str:
        """Enhanced ZIP extraction for complex submissions with multiple file types."""
        try:
            if not DocumentProcessor._verify_file(file_path):
                return ""

            all_content = ""
            file_count = 0

            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Get file list and organize by type
                file_list = zip_file.namelist()

                # Organize files by type
                text_files = []
                code_files = []
                document_files = []
                other_files = []

                for file_name in file_list:
                    if file_name.endswith('/') or file_name.startswith('__MACOSX'):
                        continue  # Skip directories and Mac metadata

                    if file_name.endswith(('.txt', '.md', '.rst')):
                        text_files.append(file_name)
                    elif file_name.endswith(('.py', '.java', '.cpp', '.c', '.h', '.js', '.html', '.css', '.sql', '.r', '.m')):
                        code_files.append(file_name)
                    elif file_name.endswith(('.pdf', '.docx', '.doc')):
                        document_files.append(file_name)
                    else:
                        other_files.append(file_name)

                all_content += f"=== ZIP FILE ANALYSIS ===\n"
                all_content += f"Archive: {Path(file_path).name}\n"
                all_content += f"Total Files: {len(file_list)}\n"
                all_content += f"Text Files: {len(text_files)}\n"
                all_content += f"Code Files: {len(code_files)}\n"
                all_content += f"Document Files: {len(document_files)}\n"
                all_content += f"Other Files: {len(other_files)}\n\n"

                # Extract text files
                if text_files:
                    all_content += "=== TEXT FILES ===\n"
                    for file_name in text_files[:10]:  # Limit to 10 files
                        try:
                            with zip_file.open(file_name) as f:
                                content = DocumentProcessor._extract_text_with_encoding(f)
                                if content:
                                    all_content += f"\n--- {file_name} ---\n"
                                    all_content += content + "\n"
                                    file_count += 1
                        except Exception as e:
                            logger.warning(f"Failed to extract {file_name}: {e}")

                # Extract code files
                if code_files:
                    all_content += "\n=== CODE FILES ===\n"
                    for file_name in code_files[:20]:  # Limit to 20 code files
                        try:
                            with zip_file.open(file_name) as f:
                                content = DocumentProcessor._extract_text_with_encoding(f)
                                if content:
                                    all_content += f"\n--- {file_name} ---\n"
                                    all_content += content + "\n"
                                    file_count += 1
                        except Exception as e:
                            logger.warning(f"Failed to extract {file_name}: {e}")

                # Extract document files
                if document_files:
                    all_content += "\n=== DOCUMENT FILES ===\n"
                    for file_name in document_files[:5]:  # Limit to 5 documents
                        try:
                            if file_name.endswith('.pdf'):
                                content = DocumentProcessor._extract_pdf_from_zip(zip_file, file_name)
                            elif file_name.endswith(('.docx', '.doc')):
                                content = DocumentProcessor._extract_docx_from_zip(zip_file, file_name)
                            else:
                                continue

                            if content:
                                all_content += f"\n--- {file_name} ---\n"
                                all_content += content + "\n"
                                file_count += 1
                        except Exception as e:
                            logger.warning(f"Failed to extract {file_name}: {e}")

                # List other files
                if other_files:
                    all_content += f"\n=== OTHER FILES ({len(other_files)}) ===\n"
                    for file_name in other_files[:10]:
                        all_content += f"- {file_name}\n"

            if file_count == 0:
                logger.error("No files extracted from ZIP archive")
                return ""

            # Analyze and structure the combined content
            structured_text = DocumentProcessor._analyze_and_structure_content(
                all_content, file_path, f"ZIP ({file_count} files)"
            )

            logger.info(f"Successfully extracted {file_count} files from ZIP archive")
            return structured_text

        except Exception as e:
            logger.error(f"Error extracting ZIP contents: {str(e)}")
            return ""

    @staticmethod
    def _extract_text_with_encoding(file_obj) -> str:
        """Extract text with multiple encoding attempts."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                file_obj.seek(0)  # Reset file pointer
                content = file_obj.read().decode(encoding)
                return content
            except UnicodeDecodeError:
                continue

        return "[Could not decode file - binary content]"

    @staticmethod
    def _extract_pdf_from_zip(zip_file, file_name: str) -> str:
        """Extract PDF content from ZIP file."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            try:
                with zip_file.open(file_name) as pdf_file:
                    temp_file.write(pdf_file.read())
                    temp_file.flush()
                    return DocumentProcessor.extract_text_from_pdf(temp_file.name)
            finally:
                try:
                    os.unlink(temp_file.name)
                except:
                    pass

    @staticmethod
    def _extract_docx_from_zip(zip_file, file_name: str) -> str:
        """Extract DOCX content from ZIP file."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            try:
                with zip_file.open(file_name) as docx_file:
                    temp_file.write(docx_file.read())
                    temp_file.flush()
                    return DocumentProcessor.extract_text_from_docx(temp_file.name)
            finally:
                try:
                    os.unlink(temp_file.name)
                except:
                    pass

    @staticmethod
    def _extract_text_file(file_path: str, original_name: str) -> str:
        """Extract text from various text file formats with multiple encoding support."""
        try:
            # Analyze file extension
            file_ext = Path(original_name).suffix.lower()

            # Common text file extensions
            text_extensions = [
                '.txt', '.md', '.rst', '.log', '.py', '.java', '.cpp', '.c', '.h',
                '.js', '.html', '.css', '.xml', '.json', '.csv', '.sql', '.r',
                '.m', '.sh', '.bat', '.ps1', '.yaml', '.yml', '.ini', '.cfg'
            ]

            content = ""

            if file_ext in text_extensions or file_ext == '':
                # Try multiple encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']

                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                            break
                    except UnicodeDecodeError:
                        continue
                    except Exception:
                        continue

                if content:
                    # Structure the content
                    structured_content = f"=== TEXT FILE ANALYSIS ===\n"
                    structured_content += f"File: {original_name}\n"
                    structured_content += f"Extension: {file_ext}\n"
                    structured_content += f"Content Length: {len(content)} characters\n\n"

                    # Detect content type
                    if file_ext in ['.py', '.java', '.cpp', '.c', '.h', '.js']:
                        structured_content += "Content Type: Source Code\n"
                    elif file_ext in ['.md', '.rst']:
                        structured_content += "Content Type: Documentation\n"
                    elif file_ext in ['.html', '.css', '.xml']:
                        structured_content += "Content Type: Markup/Style\n"
                    elif file_ext in ['.json', '.yaml', '.yml']:
                        structured_content += "Content Type: Configuration/Data\n"
                    else:
                        structured_content += "Content Type: Text Document\n"

                    structured_content += "\n=== FILE CONTENT ===\n"
                    structured_content += content

                    return structured_content
                else:
                    logger.error(f"Could not decode text file: {original_name}")
                    return ""
            else:
                logger.error(f"Unsupported file type: {file_ext}")
                return ""

        except Exception as e:
            logger.error(f"Error extracting text file: {e}")
            return ""

class RAGSystem:
    """Retrieval-Augmented Generation system for feedback generation using Claude API."""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
        self.client = None
        self.embedder = None
        self.index = None
        self.document_chunks = []
        self.rubric_text = ""
        
        if RAG_AVAILABLE:
            self._initialize_models()
    
    def _initialize_models(self):
        """Initialize the Claude API client and embedding model."""
        try:
            # Initialize embedding model
            self.embedder = SentenceTransformer('BAAI/bge-small-en-v1.5')
            logger.info("Embedding model loaded successfully")
            
            # Initialize Claude API client
            if self.api_key:
                self.client = anthropic.Anthropic(api_key=self.api_key)
                logger.info("Claude API client initialized successfully")
            else:
                logger.warning("No API key provided. Set ANTHROPIC_API_KEY environment variable or pass api_key parameter")
                
        except Exception as e:
            logger.error(f"Error initializing RAG models: {str(e)}")
    
    def load_rubric(self, rubric_text: str):
        """Load and process marking rubric for context."""
        self.rubric_text = rubric_text
        if self.embedder:
            # Create embeddings for rubric sections
            rubric_chunks = self._chunk_text(rubric_text, chunk_size=300)
            self._update_index(rubric_chunks)

    def add_document(self, file_path: str, document_name: str):
        """Add a document file to the RAG system for context - supports all CS/DS submission types."""
        try:
            file_extension = Path(file_path).suffix.lower()
            extracted_text = ""

            # Enhanced document processing for CS/DS submissions
            if file_extension == '.pdf':
                extracted_text = DocumentProcessor.extract_text_from_pdf(file_path)

            elif file_extension == '.docx':
                extracted_text = DocumentProcessor.extract_text_from_docx(file_path)

            elif file_extension == '.zip':
                # Handle ZIP archives (common for programming assignments)
                zip_contents = DocumentProcessor.extract_text_from_zip(file_path)
                extracted_text = self._process_zip_contents(zip_contents, document_name)

            elif file_extension in ['.txt', '.md', '.readme']:
                # Text and markdown files
                extracted_text = self._read_text_file(file_path)

            # Programming language files
            elif file_extension in ['.py', '.java', '.cpp', '.c', '.h', '.hpp', '.js', '.ts',
                                   '.html', '.css', '.php', '.rb', '.go', '.rs', '.swift',
                                   '.kt', '.scala', '.r', '.m', '.sh', '.bat', '.ps1']:
                extracted_text = self._process_code_file(file_path, file_extension)

            # Data files
            elif file_extension in ['.csv', '.json', '.xml', '.yaml', '.yml']:
                extracted_text = self._process_data_file(file_path, file_extension)

            # Jupyter notebooks
            elif file_extension == '.ipynb':
                extracted_text = self._process_jupyter_notebook(file_path)

            # Configuration and project files
            elif file_extension in ['.config', '.ini', '.properties', '.toml']:
                extracted_text = self._read_text_file(file_path)

            # Documentation files
            elif file_extension in ['.rst', '.tex', '.latex']:
                extracted_text = self._read_text_file(file_path)

            # Database files (basic text extraction)
            elif file_extension in ['.sql', '.sqlite', '.db']:
                extracted_text = self._process_database_file(file_path, file_extension)

            else:
                # Try reading as text file (fallback)
                extracted_text = self._read_text_file(file_path)

            # Process extracted content
            if extracted_text and extracted_text.strip():
                # Determine document type and process accordingly
                doc_type = self._classify_document_type(extracted_text, document_name, file_extension)

                if doc_type == 'rubric':
                    logger.info(f"Loading {document_name} as marking rubric")
                    self.load_rubric(extracted_text)
                else:
                    logger.info(f"Adding {document_name} as {doc_type} document")
                    # Enhanced chunking for different content types
                    doc_chunks = self._intelligent_chunk_text(extracted_text, doc_type)
                    self._update_index(doc_chunks)

                logger.info(f"Successfully processed {document_name} ({file_extension}) as {doc_type}")
            else:
                logger.warning(f"No extractable content found in {file_path}")

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise

    def _process_zip_contents(self, zip_contents: Dict[str, str], document_name: str) -> str:
        """Process ZIP archive contents for programming assignments."""
        combined_text = f"=== ZIP Archive: {document_name} ===\n\n"

        # Categorize files by type
        code_files = []
        doc_files = []
        data_files = []

        for filename, content in zip_contents.items():
            file_ext = Path(filename).suffix.lower()

            if file_ext in ['.py', '.java', '.cpp', '.c', '.h', '.js', '.html', '.css']:
                code_files.append((filename, content))
            elif file_ext in ['.txt', '.md', '.pdf', '.docx']:
                doc_files.append((filename, content))
            elif file_ext in ['.csv', '.json', '.xml']:
                data_files.append((filename, content))
            else:
                code_files.append((filename, content))  # Default to code

        # Process in logical order
        if doc_files:
            combined_text += "=== DOCUMENTATION FILES ===\n"
            for filename, content in doc_files:
                combined_text += f"\n--- {filename} ---\n{content[:2000]}\n"

        if code_files:
            combined_text += "\n=== CODE FILES ===\n"
            for filename, content in code_files:
                combined_text += f"\n--- {filename} ---\n{content[:3000]}\n"

        if data_files:
            combined_text += "\n=== DATA FILES ===\n"
            for filename, content in data_files:
                combined_text += f"\n--- {filename} ---\n{content[:1000]}\n"

        return combined_text

    def _process_code_file(self, file_path: str, file_extension: str) -> str:
        """Process programming language files with syntax awareness."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Add language context
            language_map = {
                '.py': 'Python', '.java': 'Java', '.cpp': 'C++', '.c': 'C',
                '.js': 'JavaScript', '.ts': 'TypeScript', '.html': 'HTML',
                '.css': 'CSS', '.php': 'PHP', '.rb': 'Ruby', '.go': 'Go',
                '.rs': 'Rust', '.swift': 'Swift', '.kt': 'Kotlin', '.r': 'R'
            }

            language = language_map.get(file_extension, 'Code')
            processed_content = f"=== {language} Code File ===\n\n{content}"

            return processed_content

        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f"=== Code File (Binary/Encoded) ===\n\n{f.read()}"
            except:
                return f"=== Binary File ===\nCould not extract text from {Path(file_path).name}"

    def _process_data_file(self, file_path: str, file_extension: str) -> str:
        """Process data files (CSV, JSON, XML) with structure awareness."""
        try:
            if file_extension == '.csv':
                # Sample first few rows of CSV
                df = pd.read_csv(file_path, nrows=50)
                structure = f"CSV Structure: {df.shape[0]}+ rows, {df.shape[1]} columns\n"
                structure += f"Columns: {', '.join(df.columns.tolist())}\n"
                structure += f"Sample data:\n{df.head(10).to_string()}"
                return f"=== CSV Data File ===\n\n{structure}"

            elif file_extension == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()[:5000]  # Limit JSON size
                return f"=== JSON Data File ===\n\n{content}"

            elif file_extension in ['.xml', '.yaml', '.yml']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()[:3000]
                return f"=== {file_extension.upper()} Data File ===\n\n{content}"

        except Exception as e:
            return f"=== Data File (Error) ===\nCould not process: {str(e)}"

    def _process_jupyter_notebook(self, file_path: str) -> str:
        """Process Jupyter notebook files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                notebook_content = f.read()

            # Basic notebook processing (could be enhanced with nbformat)
            import json
            notebook = json.loads(notebook_content)

            processed_content = "=== Jupyter Notebook ===\n\n"

            for i, cell in enumerate(notebook.get('cells', [])[:20]):  # Limit cells
                cell_type = cell.get('cell_type', 'unknown')
                source = ''.join(cell.get('source', []))

                if source.strip():
                    processed_content += f"--- Cell {i+1} ({cell_type}) ---\n{source}\n\n"

            return processed_content

        except Exception as e:
            return f"=== Jupyter Notebook (Error) ===\nCould not process: {str(e)}"

    def _process_database_file(self, file_path: str, file_extension: str) -> str:
        """Process database-related files."""
        if file_extension == '.sql':
            return self._read_text_file(file_path)
        else:
            return f"=== Database File ===\n{Path(file_path).name} - Binary database file"

    def _read_text_file(self, file_path: str) -> str:
        """Safely read text files with encoding fallback."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        return f"Could not decode text file: {Path(file_path).name}"

    def _classify_document_type(self, content: str, filename: str, file_extension: str) -> str:
        """Classify document type for appropriate processing."""
        content_lower = content.lower()

        # Rubric detection (enhanced)
        rubric_keywords = [
            'rubric', 'marking', 'criteria', 'grade', 'assessment', 'scoring',
            'points', 'marks', 'evaluation', 'grading scheme', 'marking scheme'
        ]

        if any(keyword in content_lower for keyword in rubric_keywords):
            return 'rubric'

        # Programming assignment
        if file_extension in ['.py', '.java', '.cpp', '.c', '.js'] or 'class ' in content or 'function ' in content:
            return 'code'

        # Data science content
        if any(keyword in content_lower for keyword in ['pandas', 'numpy', 'matplotlib', 'sklearn', 'tensorflow', 'pytorch']):
            return 'data_science'

        # Research report
        if any(keyword in content_lower for keyword in ['abstract', 'introduction', 'methodology', 'results', 'conclusion']):
            return 'report'

        # Documentation
        if file_extension in ['.md', '.rst', '.txt'] and any(keyword in content_lower for keyword in ['readme', 'documentation', 'manual']):
            return 'documentation'

        return 'general'

    def _intelligent_chunk_text(self, text: str, doc_type: str) -> List[str]:
        """Intelligent text chunking based on document type."""
        if doc_type == 'code':
            # Preserve code structure - larger chunks
            return self._chunk_text(text, chunk_size=800, overlap=100)
        elif doc_type == 'data_science':
            # Medium chunks for data science content
            return self._chunk_text(text, chunk_size=600, overlap=75)
        elif doc_type == 'report':
            # Standard chunks for reports
            return self._chunk_text(text, chunk_size=500, overlap=50)
        else:
            # Default chunking
            return self._chunk_text(text, chunk_size=500, overlap=50)
    
    def _chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks for better retrieval."""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk.strip())
        
        return chunks
    
    def _update_index(self, new_chunks: List[str]):
        """Update FAISS index with new document chunks."""
        if not self.embedder:
            return
        
        try:
            # Add new chunks to existing ones
            self.document_chunks.extend(new_chunks)
            
            # Create embeddings for all chunks
            embeddings = self.embedder.encode(self.document_chunks)
            
            # Create or update FAISS index
            dimension = embeddings.shape[1]
            if self.index is None:
                self.index = faiss.IndexFlatIP(dimension)  # Inner product similarity
            else:
                # Reset index with new data
                self.index = faiss.IndexFlatIP(dimension)
            
            # Normalize embeddings for cosine similarity
            faiss.normalize_L2(embeddings)
            self.index.add(embeddings.astype(np.float32))
            
            logger.info(f"Index updated with {len(self.document_chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Error updating index: {str(e)}")
    
    def retrieve_relevant_context(self, query: str, k: int = 5) -> List[str]:
        """Retrieve most relevant document chunks for a query."""
        if not self.embedder or not self.index:
            return []
        
        try:
            # Encode query
            query_embedding = self.embedder.encode([query])
            faiss.normalize_L2(query_embedding)
            
            # Search index
            scores, indices = self.index.search(query_embedding.astype(np.float32), k)
            
            # Return relevant chunks
            relevant_chunks = []
            for idx in indices[0]:
                if idx < len(self.document_chunks):
                    relevant_chunks.append(self.document_chunks[idx])
            
            return relevant_chunks
            
        except Exception as e:
            logger.error(f"Error retrieving context: {str(e)}")
            return []
    
    def generate_feedback(self, submission_text: str, query: str = None) -> str:
        """Generate feedback using Claude API with RAG approach."""
        if not self.client:
            return "RAG system not available - Claude API client not initialized"
        
        try:
            # Use default query if none provided
            if not query:
                query = "Evaluate this student submission against the marking rubric and provide constructive feedback"
            
            # Retrieve relevant context
            context_chunks = self.retrieve_relevant_context(query, k=3)
            context = "\n\n".join(context_chunks)
            
            # Prepare prompt
            prompt = self._create_feedback_prompt(submission_text, context, query)
            
            # Generate response using Claude API
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=1500,
                temperature=0.7,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Extract response text
            feedback = response.content[0].text.strip()
            
            # Post-process feedback
            feedback = self._post_process_feedback(feedback)
            
            logger.info("Feedback generated successfully using Claude API")
            return feedback
            
        except Exception as e:
            logger.error(f"Error generating feedback: {str(e)}")
            return f"Error generating feedback: {str(e)}"
    
    def _create_feedback_prompt(self, submission_text: str, context: str, query: str) -> str:
        """Create structured prompt for feedback generation with enhanced academic report support."""

        # Detect if this is an academic report
        is_academic_report = any(indicator in submission_text.lower() for indicator in [
            'abstract', 'introduction', 'methodology', 'results', 'conclusion',
            'figure', 'table', 'appendix', 'references'
        ])

        # Detect if contains code
        has_code = any(indicator in submission_text.lower() for indicator in [
            'def ', 'class ', 'function', 'import ', '#include', 'public class',
            'algorithm', 'code:', 'listing'
        ])

        if is_academic_report:
            prompt = f"""You are an expert academic assessor specializing in Computer Science and Data Science reports. You are evaluating a comprehensive academic report that includes written content, figures, tables, flowcharts, and code appendices.

ASSESSMENT FRAMEWORK - Evaluate ALL components:

📝 WRITTEN CONTENT ANALYSIS:
1. **Academic Structure**: Abstract, Introduction, Methodology, Results, Discussion, Conclusion
2. **Argument Quality**: Logical flow, evidence support, critical analysis
3. **Literature Integration**: Use of sources, citations, academic rigor
4. **Writing Quality**: Clarity, coherence, academic style, grammar

📊 VISUAL ELEMENTS EVALUATION:
1. **Figures & Charts**: Relevance, clarity, proper labeling, integration with text
2. **Tables & Data**: Structure, completeness, appropriate formatting
3. **Flowcharts & Diagrams**: Logic flow, clarity, technical accuracy
4. **Visual-Text Integration**: How well visuals support the narrative

💻 CODE & TECHNICAL ASSESSMENT:
1. **Code Quality**: Structure, readability, efficiency, best practices
2. **Implementation**: Correctness, completeness, innovation
3. **Documentation**: Comments, explanations, code-text integration
4. **Technical Rigor**: Algorithmic approach, problem-solving methodology

🎯 HOLISTIC EVALUATION:
1. **Coherence**: How well all components work together
2. **Academic Rigor**: Depth of analysis, methodology soundness
3. **Innovation**: Original contributions, creative solutions
4. **Professional Presentation**: Overall quality and attention to detail

FEEDBACK STRUCTURE - Provide comprehensive assessment:

**STRENGTHS** (What works well):
- Written content quality
- Visual element effectiveness
- Technical implementation
- Overall presentation

**AREAS FOR IMPROVEMENT** (Specific suggestions):
- Content development needs
- Visual enhancement opportunities
- Code optimization suggestions
- Structural improvements

**SPECIFIC RECOMMENDATIONS** (Actionable advice):
- Concrete steps for improvement
- Reference to rubric criteria
- Examples from the submission

Marking Context:
{context}

Rubric Information:
{self.rubric_text}

Student Submission Analysis:
{submission_text[:6000]}

Assessment Focus: {query}

Provide detailed, fair feedback that acknowledges the multi-faceted nature of this academic report. Address written content, visual elements, and technical components with equal attention. Be specific in your examples and constructive in your suggestions."""

        else:
            prompt = f"""You are an expert academic marker providing constructive feedback on student submissions.

Your task is to:
1. Evaluate the submission against the provided marking criteria
2. Provide specific, constructive feedback
3. Highlight both strengths and areas for improvement
4. Be encouraging but honest in your assessment
5. Reference specific parts of the submission in your feedback

Important: Your feedback is for reference only and will assist human markers in their evaluation. Always maintain academic integrity and provide balanced, educational feedback.

Marking Context:
{context}

Rubric Information:
{self.rubric_text}

Student Submission:
{submission_text[:4000]}

Query: {query}

Please provide detailed feedback for this submission, focusing on both strengths and areas for improvement. Structure your feedback clearly with specific examples from the submission."""

        return prompt
    
    def _post_process_feedback(self, feedback: str) -> str:
        """Clean and structure the generated feedback."""
        # Clean up common formatting issues
        lines = feedback.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('Human:') and not line.startswith('User:'):
                cleaned_lines.append(line)
        
        feedback = '\n'.join(cleaned_lines)
        
        # Add disclaimer
        disclaimer = "\n\n--- AI-Generated Feedback (Claude API - For Reference Only) ---\nThis feedback was generated by Claude AI to support the marking process. Human markers should use their professional judgment for final assessment decisions."
        
        return feedback + disclaimer

class FeedbackGenerator:
    """Main interface for generating AI feedback on student submissions."""
    
    def __init__(self, api_key: str = None):
        self.rag_system = RAGSystem(api_key) if RAG_AVAILABLE else None
        self.doc_processor = DocumentProcessor()
    
    def generate_feedback(self, submission_file, rubric_file=None) -> str:
        """
        Generate AI feedback for a student submission.
        
        Args:
            submission_file: Uploaded submission file (PDF, DOCX, or ZIP)
            rubric_file: Optional rubric file for context
            
        Returns:
            str: Generated feedback text
        """
        if not self.rag_system:
            return """AI Feedback System Not Available

The RAG-based feedback system requires additional dependencies:
- anthropic
- sentence-transformers  
- faiss-cpu

Please install these packages and ensure your Claude API key is set in the ANTHROPIC_API_KEY environment variable.

Alternative: Provide manual feedback based on the submission and rubric."""
        
        try:
            # Extract text from submission
            submission_text = self._extract_submission_text(submission_file)
            
            if not submission_text.strip():
                return "Error: Could not extract text from the submission file."
            
            # Process rubric if provided
            if rubric_file:
                rubric_text = self._extract_rubric_text(rubric_file)
                if rubric_text:
                    self.rag_system.load_rubric(rubric_text)
            
            # Update RAG index with submission content
            submission_chunks = self.rag_system._chunk_text(submission_text)
            self.rag_system._update_index(submission_chunks)
            
            # Generate feedback
            feedback = self.rag_system.generate_feedback(
                submission_text,
                "Provide comprehensive feedback on this student submission, highlighting strengths and areas for improvement"
            )
            
            return feedback
            
        except Exception as e:
            logger.error(f"Error in feedback generation: {str(e)}")
            return f"Error generating feedback: {str(e)}"
    
    def _extract_submission_text(self, submission_file) -> str:
        """Extract text from uploaded submission file with enhanced debugging."""
        try:
            logger.info(f"Processing file: {submission_file.name}")

            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(submission_file.name).suffix) as temp_file:
                temp_file.write(submission_file.getvalue())
                temp_file.flush()

                logger.info(f"Temporary file created: {temp_file.name}")

                # Extract text based on file type
                file_extension = Path(submission_file.name).suffix.lower()
                logger.info(f"File extension detected: {file_extension}")

                text = ""

                if file_extension == '.pdf':
                    logger.info("Using PDF extraction")
                    text = self.doc_processor.extract_text_from_pdf(temp_file.name)
                elif file_extension == '.docx':
                    logger.info("Using DOCX extraction")
                    text = self.doc_processor.extract_text_from_docx(temp_file.name)
                elif file_extension == '.zip':
                    logger.info("Using ZIP extraction")
                    text = self.doc_processor.extract_text_from_zip(temp_file.name)
                else:
                    logger.info("Using text file extraction")
                    text = DocumentProcessor._extract_text_file(temp_file.name, submission_file.name)

                # Clean up temporary file
                try:
                    os.unlink(temp_file.name)
                except:
                    pass

                logger.info(f"Extracted text length: {len(text) if text else 0} characters")

                if not text or not text.strip():
                    logger.error("No text extracted from file")
                    return ""

                return text

        except Exception as e:
            logger.error(f"Error extracting submission text: {str(e)}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return ""
    
    def _extract_rubric_text(self, rubric_file) -> str:
        """Extract text from rubric file."""
        try:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(rubric_file.name).suffix) as temp_file:
                temp_file.write(rubric_file.getvalue())
                temp_file.flush()
                
                # Extract text based on file type
                file_extension = Path(rubric_file.name).suffix.lower()
                
                if file_extension == '.pdf':
                    text = self.doc_processor.extract_text_from_pdf(temp_file.name)
                elif file_extension == '.docx':
                    text = self.doc_processor.extract_text_from_docx(temp_file.name)
                else:
                    # Try to read as text file with multiple encodings
                    text = DocumentProcessor._extract_text_file(temp_file.name, rubric_file.name)
                
                # Clean up temporary file
                os.unlink(temp_file.name)
                
                return text
                
        except Exception as e:
            logger.error(f"Error extracting rubric text: {str(e)}")
            return ""
    
    def test_system(self) -> Dict[str, Any]:
        """Test the feedback generation system."""
        test_results = {
            'rag_available': RAG_AVAILABLE,
            'model_loaded': False,
            'embedder_loaded': False,
            'test_feedback': None,
            'errors': []
        }
        
        if not RAG_AVAILABLE:
            test_results['errors'].append("RAG dependencies not installed")
            return test_results
        
        try:
            # Test Claude API client loading
            if self.rag_system and self.rag_system.client:
                test_results['model_loaded'] = True
            
            if self.rag_system and self.rag_system.embedder:
                test_results['embedder_loaded'] = True
            
            # Generate test feedback
            if self.rag_system and self.rag_system.client:
                test_text = "This is a simple test submission to verify the feedback system is working correctly."
                test_feedback = self.rag_system.generate_feedback(test_text, "Provide brief test feedback")
                test_results['test_feedback'] = test_feedback
            
        except Exception as e:
            test_results['errors'].append(str(e))
        
        return test_results